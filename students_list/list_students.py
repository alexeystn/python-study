from docx import Document
import os


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)



path = '/Users/imac/Русский язык/Отчёты о количестве учащихся/'

files = os.listdir(path)

students = {}


for f in files:
    if f.endswith('docx') & (len(f) == 15):
        print(f)
        doc = Document(path + f)
        for row in doc.tables[0].rows[1:]:
            st_name = row.cells[0].text.strip().replace('ё', 'е')
            st_birth = row.cells[1].text.strip()
            st_class = row.cells[2].text.strip()
            st_city = row.cells[3].text.strip()
            st_school = row.cells[4].text.strip()
            st_arrival = row.cells[5].text.strip()

            entry = [st_birth, st_class, st_city, st_school]

            if not st_name in students.keys():
                students[st_name] = {}

            if not st_arrival in students[st_name].keys():
                students[st_name][st_arrival] = entry

student_names = list(students.keys())
student_names.sort()

f = open('students.txt', 'w', encoding = 'utf-8')
for s in student_names:
    for i, arrival_date in enumerate(students[s].keys()):
        if i == 0:
            f.write(s)
        f.write('\t')
        f.write(students[s][arrival_date][0] + '\t') # birth
        f.write(students[s][arrival_date][1] + '\t')  # class
        f.write(students[s][arrival_date][2] + '\t')  # city
        f.write(students[s][arrival_date][3] + '\t')  # school
        f.write(arrival_date)        
        f.write('\n')

f.close()





##path = 'output.docx'
##doc = Document(path)
##
##table = doc.tables[0]
##for i in reversed(range(1, len(table.rows))):
##    row = doc.tables[0].rows[i]
##    remove_row(table, row)
##
##for s in student_names:
##    for i, arrival_date in enumerate(students[s].keys()):
##        row = doc.tables[0].add_row()
##        if i == 0:
##            row.cells[0].text = s
##        row.cells[1].text = students[s][arrival_date][0] # birth
##        row.cells[2].text = students[s][arrival_date][1] # class
##        row.cells[3].text = students[s][arrival_date][2] # city
##        row.cells[4].text = students[s][arrival_date][3] # school
##        row.cells[5].text = arrival_date
##doc.save(path)








